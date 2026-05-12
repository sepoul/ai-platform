from io import BytesIO
from typing import ClassVar, List

from docx import Document

from ai_platform.utilities.input_parsers.base import DocumentParser, ParsedDocument


class DocxDocumentParser(DocumentParser):
    """Parse .docx files using python-docx."""

    supported_content_types: ClassVar[List[str]] = [
        "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    ]

    def parse(self, file_bytes: bytes, source_name: str) -> ParsedDocument:
        doc = Document(BytesIO(file_bytes))

        parts: List[str] = []

        # Paragraphs
        for p in doc.paragraphs:
            text = p.text.strip()
            if text:
                parts.append(text)

        # Tables
        for table in doc.tables:
            for row in table.rows:
                row_text = "\t".join(cell.text.strip() for cell in row.cells)
                if row_text.strip():
                    parts.append(row_text)

        return ParsedDocument(
            source_name=source_name,
            content_type=self.supported_content_types[0],
            text="\n".join(parts),
            page_count=None,  # python-docx doesn't expose page count
        )
